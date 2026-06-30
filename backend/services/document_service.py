from __future__ import annotations

import csv
import hashlib
import io
import json
import logging
import re
import time
from datetime import datetime, timezone
from pathlib import PurePosixPath
from typing import TYPE_CHECKING, Any
from uuid import uuid4

from fastapi import UploadFile

from config.settings import Settings
from providers.storage_provider import StorageProvider
from providers.supabase_provider import SupabaseProvider
from schemas.document import Document, DocumentStatistics, DocumentStatus
from utils.exceptions import AppError

if TYPE_CHECKING:
    from providers.redis_provider import RedisProvider
    from services.llm_service import LLMService

logger = logging.getLogger(__name__)


SUPPORTED_FORMATS = {"pdf", "md", "txt", "docx", "csv", "json", "png", "jpg", "jpeg", "webp", "mp3", "wav", "m4a"}
IMAGE_FORMATS = {"png", "jpg", "jpeg", "webp"}
AUDIO_FORMATS = {"mp3", "wav", "m4a"}


class OCRProvider:
    async def extract_text(self, _: bytes, filename: str) -> str:
        return f"OCR is pending for image document: {filename}"


class TranscriptionProvider:
    async def transcribe(self, _: bytes, filename: str) -> str:
        return f"Transcription is pending for audio document: {filename}"


class DocumentService:
    def __init__(
        self,
        settings: Settings,
        supabase_provider: SupabaseProvider,
        storage_provider: StorageProvider,
        llm_service: LLMService,
        redis_provider: RedisProvider | None = None,
        ocr_provider: OCRProvider | None = None,
        transcription_provider: TranscriptionProvider | None = None,
    ) -> None:
        self.settings = settings
        self.supabase_provider = supabase_provider
        self.storage_provider = storage_provider
        self.llm_service = llm_service
        self.redis_provider = redis_provider
        self.ocr_provider = ocr_provider or OCRProvider()
        self.transcription_provider = transcription_provider or TranscriptionProvider()

    async def upload_document(self, user_id: str, file: UploadFile) -> Document:
        started_at = time.perf_counter()
        content = await file.read()
        filename = file.filename or "document"
        file_type = self._file_type(filename)
        file_size = len(content)

        self._validate_upload(filename, file_type, file_size)
        await self._ensure_not_duplicate(user_id, filename, file_size)

        document_id = str(uuid4())
        storage_path = self._storage_path(user_id, document_id, filename)
        storage_url = await self.storage_provider.upload(storage_path, content, file.content_type)
        now = datetime.now(timezone.utc).isoformat()

        record = await self.supabase_provider.insert_one(
            "documents",
            {
                "id": document_id,
                "filename": filename,
                "file_type": file_type,
                "file_size": file_size,
                "created_at": now,
                "uploaded_by": user_id,
                "summary": None,
                "keywords": [],
                "language": None,
                "reading_time": None,
                "status": DocumentStatus.pending.value,
                "storage_url": storage_url,
                "storage_path": storage_path,
                "metadata": {"progress": 0, "content_hash": self._content_hash(content)},
            },
        )
        if not record:
            await self.storage_provider.delete(storage_path)
            raise AppError(message="Failed to create document record", error="document_create_failed", status_code=502)

        document = self._to_document(record)
        await self._cache_document(document)
        logger.info(
            "document_uploaded user_id=%s filename=%s size=%s processing_time=%.3fs",
            user_id,
            filename,
            file_size,
            time.perf_counter() - started_at,
        )
        return document

    async def list_documents(self, user_id: str) -> list[Document]:
        records = await self.supabase_provider.select_many("documents", {"uploaded_by": user_id}, order_by="created_at")
        documents = [self._to_document(record) for record in records]
        return sorted(documents, key=lambda document: document.created_at or datetime.min, reverse=True)

    async def get_document(self, user_id: str, document_id: str) -> Document:
        cached = await self._get_cached_document(document_id)
        if cached and cached.uploaded_by == user_id:
            return cached

        record = await self.supabase_provider.select_one("documents", {"id": document_id, "uploaded_by": user_id})
        if not record:
            raise AppError(message="Document not found", error="document_not_found", status_code=404)
        document = self._to_document(record)
        await self._cache_document(document)
        return document

    async def delete_document(self, user_id: str, document_id: str) -> None:
        record = await self.supabase_provider.delete_one("documents", {"id": document_id, "uploaded_by": user_id})
        if not record:
            raise AppError(message="Document not found", error="document_not_found", status_code=404)

        storage_path = record.get("storage_path")
        if storage_path:
            await self.storage_provider.delete(str(storage_path))
        await self._delete_cached_document(document_id)

    async def process_document(self, user_id: str, document_id: str, content: bytes) -> None:
        started_at = time.perf_counter()
        document = await self.get_document(user_id, document_id)
        await self._update_document(document_id, {"status": DocumentStatus.processing.value, "metadata": {**document.metadata, "progress": 20}})

        try:
            text = await self.extract_text(content, document.filename, document.file_type)
            await self._update_document(document_id, {"metadata": {**document.metadata, "progress": 45, "text_length": len(text)}})
            summary = await self.generate_summary(text)
            keywords = await self.extract_keywords(text)
            title = await self._generate_title(text, document.filename)
            language = self._detect_language(text)
            reading_time = self._estimated_reading_time(text)
            metadata = {
                **document.metadata,
                "progress": 100,
                "suggested_title": title,
                "text_length": len(text),
                "processed_at": datetime.now(timezone.utc).isoformat(),
            }
            updated = await self._update_document(
                document_id,
                {
                    "summary": summary,
                    "keywords": keywords,
                    "language": language,
                    "reading_time": reading_time,
                    "status": DocumentStatus.completed.value,
                    "metadata": metadata,
                },
            )
            await self._store_temporary_knowledge(user_id, document_id, text, summary, keywords)
            logger.info(
                "document_processed user_id=%s filename=%s size=%s processing_time=%.3fs",
                user_id,
                document.filename,
                document.file_size,
                time.perf_counter() - started_at,
            )
            if updated:
                await self._cache_document(self._to_document(updated))
        except Exception as exc:
            await self._update_document(
                document_id,
                {"status": DocumentStatus.failed.value, "metadata": {**document.metadata, "progress": 100, "error": exc.__class__.__name__}},
            )
            logger.exception("document_processing_failed user_id=%s filename=%s", user_id, document.filename)

    async def extract_text(self, content: bytes, filename: str, file_type: str) -> str:
        try:
            if file_type == "pdf":
                return self._extract_pdf(content)
            if file_type == "docx":
                return self._extract_docx(content)
            if file_type == "csv":
                return self._extract_csv(content)
            if file_type == "json":
                return json.dumps(json.loads(content.decode("utf-8")), indent=2)
            if file_type in {"md", "txt"}:
                return content.decode("utf-8")
            if file_type in IMAGE_FORMATS:
                return await self.ocr_provider.extract_text(content, filename)
            if file_type in AUDIO_FORMATS:
                return await self.transcription_provider.transcribe(content, filename)
        except UnicodeDecodeError as exc:
            raise AppError(message="File must be valid UTF-8 text", error="document_text_decode_failed", status_code=400) from exc
        except Exception as exc:
            raise AppError(message="Failed to extract document text", error="document_text_extraction_failed", status_code=400) from exc

        raise AppError(message="Unsupported document format", error="unsupported_document_format", status_code=400)

    async def generate_summary(self, text: str) -> str:
        if not text.strip():
            return ""
        return (await self.llm_service.summarize(self._clip_for_llm(text))).summary

    async def extract_keywords(self, text: str) -> list[str]:
        if not text.strip():
            return []
        return (await self.llm_service.extract_keywords(self._clip_for_llm(text))).keywords

    async def document_statistics(self, user_id: str) -> DocumentStatistics:
        documents = await self.list_documents(user_id)
        return DocumentStatistics(
            total_documents=len(documents),
            total_size=sum(document.file_size for document in documents),
            supported_formats=sorted(SUPPORTED_FORMATS),
            recent_uploads=documents[:5],
        )

    async def _generate_title(self, text: str, fallback: str) -> str:
        if not text.strip():
            return fallback
        return (await self.llm_service.generate_title(self._clip_for_llm(text))).title

    async def _update_document(self, document_id: str, data: dict[str, object]) -> dict[str, object] | None:
        updated = await self.supabase_provider.update_one("documents", {"id": document_id}, data)
        await self._delete_cached_document(document_id)
        return updated

    async def _store_temporary_knowledge(
        self,
        user_id: str,
        document_id: str,
        text: str,
        summary: str,
        keywords: list[str],
    ) -> None:
        try:
            await self.supabase_provider.insert_one(
                "document_knowledge",
                {
                    "id": str(uuid4()),
                    "document_id": document_id,
                    "uploaded_by": user_id,
                    "content": text,
                    "summary": summary,
                    "keywords": keywords,
                    "created_at": datetime.now(timezone.utc).isoformat(),
                },
            )
        except Exception:
            logger.warning("temporary_document_knowledge_store_failed document_id=%s", document_id, exc_info=True)

    def _extract_pdf(self, content: bytes) -> str:
        import fitz

        with fitz.open(stream=content, filetype="pdf") as document:
            return "\n".join(page.get_text("text") for page in document)

    def _extract_docx(self, content: bytes) -> str:
        from docx import Document as DocxDocument

        document = DocxDocument(io.BytesIO(content))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)

    def _extract_csv(self, content: bytes) -> str:
        rows = csv.reader(io.StringIO(content.decode("utf-8")))
        return "\n".join(", ".join(row) for row in rows)

    def _validate_upload(self, filename: str, file_type: str, file_size: int) -> None:
        if file_size == 0:
            raise AppError(message="Uploaded file is empty", error="empty_document", status_code=400)
        if file_type not in SUPPORTED_FORMATS:
            raise AppError(message="Unsupported document format", error="unsupported_document_format", status_code=400)
        max_bytes = self.settings.document_max_file_size_mb * 1024 * 1024
        if file_size > max_bytes:
            raise AppError(message="Uploaded file is too large", error="document_too_large", status_code=413)
        if not filename.strip():
            raise AppError(message="Filename is required", error="document_filename_required", status_code=400)

    async def _ensure_not_duplicate(self, user_id: str, filename: str, file_size: int) -> None:
        existing = await self.supabase_provider.select_many("documents", {"uploaded_by": user_id, "filename": filename})
        if any(int(record.get("file_size") or 0) == file_size for record in existing):
            raise AppError(message="Duplicate document upload", error="duplicate_document", status_code=409)

    def _to_document(self, data: dict[str, Any]) -> Document:
        metadata = data.get("metadata") if isinstance(data.get("metadata"), dict) else {}
        return Document(
            id=str(data.get("id")),
            filename=str(data.get("filename")),
            file_type=str(data.get("file_type")),
            file_size=int(data.get("file_size") or 0),
            created_at=self._parse_datetime(data.get("created_at")),
            uploaded_by=str(data.get("uploaded_by")),
            summary=data.get("summary"),
            keywords=list(data.get("keywords") or []),
            language=data.get("language"),
            reading_time=data.get("reading_time"),
            status=DocumentStatus(str(data.get("status") or DocumentStatus.pending.value)),
            storage_url=str(data.get("storage_url")),
            suggested_title=data.get("suggested_title") or metadata.get("suggested_title"),
            progress=int(metadata.get("progress") or 0),
            metadata=metadata,
        )

    async def _cache_document(self, document: Document) -> None:
        if self.redis_provider and self.redis_provider.client:
            await self.redis_provider.client.setex(
                self._cache_key(document.id),
                self.settings.document_cache_ttl_seconds,
                document.model_dump_json(),
            )

    async def _get_cached_document(self, document_id: str) -> Document | None:
        if not self.redis_provider or not self.redis_provider.client:
            return None
        payload = await self.redis_provider.client.get(self._cache_key(document_id))
        return Document.model_validate_json(payload) if payload else None

    async def _delete_cached_document(self, document_id: str) -> None:
        if self.redis_provider and self.redis_provider.client:
            await self.redis_provider.client.delete(self._cache_key(document_id))

    def _cache_key(self, document_id: str) -> str:
        return f"documents:{document_id}"

    def _storage_path(self, user_id: str, document_id: str, filename: str) -> str:
        safe_filename = re.sub(r"[^A-Za-z0-9._-]+", "_", PurePosixPath(filename).name).strip("._") or "document"
        return f"{user_id}/{document_id}/{safe_filename}"

    def _file_type(self, filename: str) -> str:
        return PurePosixPath(filename).suffix.lower().lstrip(".")

    def _content_hash(self, content: bytes) -> str:
        return hashlib.sha256(content).hexdigest()

    def _clip_for_llm(self, text: str) -> str:
        return text[:16000]

    def _estimated_reading_time(self, text: str) -> int:
        words = len(re.findall(r"\w+", text))
        return max(1, round(words / 220)) if words else 0

    def _detect_language(self, text: str) -> str:
        return "en" if text.strip() else "unknown"

    def _parse_datetime(self, value: Any) -> datetime | None:
        if value is None or isinstance(value, datetime):
            return value
        if isinstance(value, str):
            try:
                return datetime.fromisoformat(value.replace("Z", "+00:00"))
            except ValueError:
                return None
        return None
